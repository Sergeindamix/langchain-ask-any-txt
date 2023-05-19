import streamlit as st
from dotenv import load_dotenv
import pickle
from PyPDF2 import PdfReader
from streamlit_extras.add_vertical_space import add_vertical_space
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.llms import OpenAI
from langchain.chains.question_answering import load_qa_chain
from langchain.callbacks import get_openai_callback
import os

from youtube_transcript_api import YouTubeTranscriptApi
import re
from langchain.text_splitter import CharacterTextSplitter
from langchain import memory
from langchain.chains.conversation.base import ConversationChain
from langchain.chat_models import ChatOpenAI 
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory, ConversationEntityMemory
from langchain.memory.prompt import ENTITY_MEMORY_CONVERSATION_TEMPLATE

import textract
import requests
import io
from pydub import AudioSegment
from IPython.display import Audio, clear_output
import tempfile
import IPython
import IPython.display as ipd
from elevenlabs import generate, play, set_api_key, voices, Models
from pydub.playback import play as play_audio
from io import BytesIO
import base64
import docx2txt

from LAVIS import imgx
from agentes import promptx
from embds import embds
from multiAsk import run_comparison



st.set_page_config(page_title="🦜🔗 Ask YouTube or Docs💬")
st.header("🦜🔗 Ask YouTube or Docs💬")


def main():
    # Create a toggle widget to generate image
    show_text = st.checkbox("Generar imagen?", value=False)

    # Display some text if the toggle is on
    if show_text:
        # get user input
        input_text = ""
        question = st.text_input("What is the main topic of this text?")

        # Invocar la función y obtener los resultados
        results = run_comparison(question)

        # Verificar si los resultados son None antes de iterar
        if results is not None:
            # Mostrar los resultados en Streamlit
            for result in results:
                model_name = result['model_name']
                response = result['response']

                # Mostrar el nombre del modelo y la respuesta generada
                st.subheader(f"Modelo: {model_name}")
                st.write(f"Respuesta: {response}")
                st.write("---")  # Separador entre cada respuesta
        else:
            st.write("No se encontraron resultados.")
        
        
        promptx()
        img_path = "1.jpg"
        imgx(img_path)

from transformers.tools import HfAgent
agent = HfAgent("https://api-inference.huggingface.co/models/bigcode/starcoder")
text = "The main topic of this text is the benefits of exercise for overall health and well-being. Studies have shown that regular physical activity can help reduce the risk of chronic diseases such as heart disease, diabetes, and cancer, as well as improve mental health and cognitive function."
    
prompts = ['"Answer the question in the variable `question` about the image stored in the variable `image`. The question is in Spanish."', 
'"Identify the oldest person in the `document` and create an image showcasing the result."', 
'"Generate an image using the text given in the variable `caption`."', 
'f"Summarize the text given in {text} and read it out loud."', 
'"Answer the question in the variable `question` about the text in the variable `text`. Use the answer to generate an image."', 
'"Caption the following `image`."', 
'"<<prompt>>"']
#prompt = st.selectbox("Selecciona un prompt:", prompts)

def download_file(file_path):
    with open(file_path, 'rb') as f:
        data = f.read()
    b64 = base64.b64encode(data).decode('utf-8')
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{file_path}">Descargar archivo</a>'
    return href

def downloadDoc(user_question, response):
  import docx
      
  # Crear nuevo documento de Word
  document = docx.Document()

  # Agregar título
  document.add_heading("Respuesta", level=1)

  # Agregar subtítulo
  document.add_heading("Pregunta del usuario:", level=2)
  document.add_paragraph(user_question)

  # Agregar subtítulo
  document.add_heading("Respuesta:", level=2)
  document.add_paragraph(response)

  # Guardar documento en un archivo
  document.save("respuesta_0.docx")

  

  counter = 0

  while os.path.exists(f"respuesta_{counter}.docx"):
      counter += 1

  document.save(f"respuesta_{counter}.docx")

  st.write(download_file(f"respuesta_{counter}.docx"), unsafe_allow_html=True)
      






def askYT():
    load_dotenv()
    st.header("🦜🔗 YouTube GPT💬")
    url = st.text_input("Ingresa link de YouTube, ejemplo: https://www.youtube.com/watch?v=KczJNtexinY")
    
    # extract video ID using regular expression
    match = re.search(r"v=(\w+)", url)
    # asigna la variable 'video_id' con el id del video de YouTube
    video_id = None

    if "youtu.be/" in url:
        # extract video id from "https://youtu.be/video_id" format
        video_id = url.split("youtu.be/")[-1]
    elif "watch?v=" in url:
        # extract video id from "https://www.youtube.com/watch?v=video_id" format
        video_id = url.split("watch?v=")[-1]

    # asigna la variable 'language_code' con el código de idioma deseado
    idioma = st.selectbox("Selecciona el idioma de los subtitulos, si no existen mostrará error", ["en", "es", "fr"])

    # obtiene el transcript en el idioma deseado
    srt = YouTubeTranscriptApi.get_transcript(video_id, languages=[idioma])
    # eliminar los caracteres de formato
    # extraer solo el texto de los subtítulos
    text = ""
    for subtitle in srt:
        text += subtitle['text'] + " "
    
        
    # eliminar los caracteres de formato
    text = text.replace('\n', ' ').replace('\r', '')
    
    # Lee el archivo de subtítulos
    with open('subtitles.txt', 'w', encoding='utf-8') as file:
      # escribir cada línea de subtítulos en el archivo
      for line in srt:
          text = line['text']
          file.write(text + '\n')

    if os.path.exists('subtitles.txt'):
        text = textract.process('subtitles.txt').decode('utf-8')
        os.remove('subtitles.txt')
    else:
        # handle file not found error
        ...


    # Define una función para dividir el texto en trozos de 2048 tokens
    def split_text(text):
      # Divide el texto en trozos de 512 tokens
      tokens = text.split()
      chunks = []
      for i in range(0, len(tokens), 512):
          chunk = ' '.join(tokens[i:i+512])
          chunks.append(chunk)
      return chunks

        

    # split into chunks    
    chunks = split_text(text)
    
    store_name = "datos"
    st.write(f'{store_name}')
    st.write(chunks)
    
    # create embeddings
    if os.path.exists(f"{store_name}.pkl"):
        with open(f"{store_name}.pkl", "rb") as f:
            VectorStore = pickle.load(f)
        # st.write('Embeddings Loaded from the Disk')s
    else:
        embeddings = OpenAIEmbeddings()
        VectorStore = FAISS.from_texts(chunks, embedding=embeddings)
        with open(f"{store_name}.pkl", "wb") as f:
            pickle.dump(VectorStore, f)
        #checar si funciona
        knowledge_base = VectorStore

    # show user input
    user_question = st.text_input("Ask a question about YouTube VIDEO:")
    if user_question:    
      docs = VectorStore.similarity_search(query=user_question, k=3)
      
      llm = OpenAI()
      chain = load_qa_chain(llm, chain_type="stuff")
      with get_openai_callback() as cb:
      
        response = chain.run(input_documents=docs, question=user_question)
        print(cb)

      downloadDoc(user_question, response)

      st.write(response)
      
import transformers
from transformers import pipeline

def is_huggingface_langchain(question, text):
    # initialize the question-answering pipeline
    qa_pipeline = pipeline("question-answering", model="distilbert-base-cased-distilled-squad")

    # find the answer to the question in the text
    result = qa_pipeline(question=question, context=text)

    # check if the answer contains the words "Hugging Face" and "LangChain"
    answer = result["answer"]
    if "Hugging Face" in answer and "LangChain" in answer:
        st.write(result)
        return True        
    else:
        st.write("error")
        return answer
      
      
def txts():
    load_dotenv()
    #st.set_page_config(page_title="Ask your PDF")
    st.header("[pdf, txt, docx] 💬")
    
    # upload file
    uploaded_file = st.file_uploader("Upload your Document", type=["pdf", "docx", "txt"])
    
    # extract the text
    if uploaded_file is not None:
        if uploaded_file.type == 'application/pdf':
            with open('uploaded_file.pdf', 'wb') as f:
                f.write(uploaded_file.read())
            if os.path.exists('uploaded_file.pdf'):
                text = textract.process('uploaded_file.pdf').decode('utf-8')
                os.remove('uploaded_file.pdf')
            else:
                # handle file not found error
                ...
        elif uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            with open('uploaded_file.docx', 'wb') as f:
                f.write(uploaded_file.read())
            if os.path.exists('uploaded_file.docx'):
                text = docx2txt.process('uploaded_file.docx')
                os.remove('uploaded_file.docx')
        elif uploaded_file.type == 'text/plain':
            # handle non-PDF files            
            with open('uploaded_file.txt', 'wb') as f:
                f.write(uploaded_file.read())
            if os.path.exists('uploaded_file.txt'):
                text = textract.process('uploaded_file.txt').decode('utf-8')
                os.remove('uploaded_file.txt')

                ...
        
        # split into chunks
        text_splitter = CharacterTextSplitter(
            separator="\n",
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        chunks = text_splitter.split_text(text)
        # split into chunks    
        
        
        store_name = uploaded_file.name[:-4]
        st.write(f'{store_name}')
        st.write(chunks)
        
        # create embeddings
        if os.path.exists(f"{store_name}.pkl"):
            with open(f"{store_name}.pkl", "rb") as f:
                VectorStore = pickle.load(f)
            # st.write('Embeddings Loaded from the Disk')s
        else:
            embeddings = OpenAIEmbeddings()
            VectorStore = FAISS.from_texts(chunks, embedding=embeddings)
            with open(f"{store_name}.pkl", "wb") as f:
                pickle.dump(VectorStore, f)
            #checar si funciona
            knowledge_base = VectorStore
       
        
        # show user input
        user_question = st.text_input("Ask a question about your PDF:")
        if user_question:
            docs = knowledge_base.similarity_search(user_question)
            
            llm = OpenAI()
            chain = load_qa_chain(llm, chain_type="stuff")
            with get_openai_callback() as cb:
                response = chain.run(input_documents=docs, question=user_question)
                print(cb)
                
            st.write(response)
            downloadDoc(user_question, response)









   


  # Módulo: Crear imagen
def create_image_module():
    from PIL import Image
    from sd import sd
    st.header("Crear imagen")
    # Agrega aquí el código para crear la imagen
    question = ""
    sd(question)

# Módulo: Cargar documentos
def load_documents_module():
    st.header("Cargar documentos")
    # Agrega aquí el código para cargar documentos
    txts()

# Módulo: Cargar URL
def load_url_module():
    st.header("Cargar URL")
    # Agrega aquí el código para cargar URL    
    url = st.text_input("https://en.wikipedia.org/wiki/Cristiano_Ronaldo")
    if st.button("Generate embeddings"):
      from embeddings import run_question_answering  
      st.session_state.vector_store = run_question_answering(url)

    # Retrieve the value of vector_store from session state
    if "vector_store" in st.session_state:
        vector_store = st.session_state.vector_store
 
        # show user input
        question = st.text_input("Ask a question about URL:")
        if question:
          from preguntame import answering
          answer = answering(vector_store, question)
          st.write(answer)

    if st.button("show docs"):
      #Check if the 'source_documents' directory exists
      source_documents_dir = "/content/LAVIS/source_documents"
      if not os.path.exists(source_documents_dir):
          os.makedirs(source_documents_dir)
          st.write(f"Created directory: {source_documents_dir}")

      # ...

      # Use 'source_documents_dir' in the file operations
      context_files = [f for f in os.listdir(source_documents_dir) if os.path.isfile(os.path.join(source_documents_dir, f))]

      # ...

      context_files.sort()  # Ordenar los archivos alfabéticamente

      # Mostrar la lista de archivos de contexto disponibles
      st.sidebar.title("Archivos de Contexto")
      selected_file = st.sidebar.selectbox("Selecciona un archivo de contexto", context_files)

      # Leer el contenido del archivo de contexto seleccionado
      with open(os.path.join("/content/LAVIS/source_documents", selected_file), "r", encoding="utf-8") as file:
          context = file.read()

      st.subheader("Contexto seleccionado:")
      st.write(context)
      #response = is_huggingface_langchain(question, text)
      #st.write(response)

# Módulo: Hacer preguntas sobre YouTube
def ask_youtube_module():
    st.header("Hacer preguntas sobre YouTube")
    # Agrega aquí el código para hacer preguntas sobre YouTube
    askYT()

# Módulo: Ask text
def ask_txt_module():
    st.header("Hacer preguntas sobre Texto")
    question = st.text_input("What is the main topic of this text?")
    text = st.text_input("The main topic of this text is the benefits of exercise for overall health and well-being. Studies have shown that regular physical activity can help reduce the risk of chronic diseases such as heart disease, diabetes, and cancer, as well as improve mental health and cognitive function.")
    response = is_huggingface_langchain(question, text)
    st.write(response)

# Sidebar contents
with st.sidebar:
    st.title('🤗💬 LLM Chat App')
    st.markdown('''
    ## Acerca de:
    Esta aplicación es un chatbot impulsado por LLM construido utilizando:
    - [Streamlit](https://streamlit.io/)
    - [LangChain](https://python.langchain.com/)
    - [OpenAI](https://platform.openai.com/docs/models) LLM model
    ''')

    # Lista desplegable para seleccionar el módulo
    selected_module = st.selectbox('Seleccionar módulo', ['Crear imagen', 'Cargar documentos', 'Cargar URL', 'Hacer preguntas sobre YouTube', 'Hacer preguntas sobre Texto'])

# Contenido del módulo seleccionado
if selected_module == 'Crear imagen':
    create_image_module()
elif selected_module == 'Cargar documentos':
    load_documents_module()
elif selected_module == 'Cargar URL':
    load_url_module()
elif selected_module == 'Hacer preguntas sobre YouTube':
    ask_youtube_module()
elif selected_module == 'Hacer preguntas sobre Texto':
    ask_txt_module()

load_dotenv()

if __name__ == '__main__':
    main()
